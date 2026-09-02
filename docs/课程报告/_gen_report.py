# -*- coding: utf-8 -*-
"""按《高级系统编程技术课程报告》模板格式生成 Word 课程报告。"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HEI = '黑体'
SONG = '宋体'
TIMES = 'Times New Roman'


def set_run_font(run, ascii_font=TIMES, ea=SONG, size=Pt(12), bold=False):
    run.font.name = ascii_font
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:eastAsia'), ea)


def set_spacing_lines(par, before_lines=None, after_lines=None):
    """以“行”为单位设置段前/段后间距（Word 的 beforeLines/afterLines）。"""
    pPr = par._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    if before_lines is not None:
        spacing.set(qn('w:beforeLines'), str(before_lines))
        spacing.set(qn('w:before'), '0')
    if after_lines is not None:
        spacing.set(qn('w:afterLines'), str(after_lines))
        spacing.set(qn('w:after'), '0')


doc = Document()

# ---------- 页面设置：A4 ----------
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin, sec.bottom_margin = Cm(2.54), Cm(2.54)
sec.left_margin, sec.right_margin = Cm(3.0), Cm(2.6)

# ---------- Normal 样式：宋体小四 + Times New Roman，1.5 倍行距 ----------
normal = doc.styles['Normal']
normal.font.name = TIMES
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), SONG)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)

# ---------- 页脚页码 ----------
footer_p = sec.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run()
fld1, fld2, fld3 = OxmlElement('w:fldChar'), OxmlElement('w:instrText'), OxmlElement('w:fldChar')
fld1.set(qn('w:fldCharType'), 'begin')
fld2.set(qn('xml:space'), 'preserve')
fld2.text = ' PAGE '
fld3.set(qn('w:fldCharType'), 'end')
run._element.append(fld1)
run._element.append(fld2)
run._element.append(fld3)
set_run_font(run, size=Pt(10.5))


# ---------- 段落工厂 ----------
def chapter(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    # 章标题：黑体小二，段前 0.8 行、段后 0.5 行
    set_spacing_lines(p, 80, 50)
    set_run_font(p.add_run(text), ea=HEI, size=Pt(18), bold=False)
    p.style = doc.styles['Heading 1']
    set_run_font(p.runs[0], ea=HEI, size=Pt(18), bold=False)
    return p


def h2(text):
    p = doc.add_paragraph(style=doc.styles['Heading 2'])
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    set_spacing_lines(p, 50, 30)
    set_run_font(p.add_run(text), ea=HEI, size=Pt(14), bold=False)
    return p


def h3(text):
    p = doc.add_paragraph(style=doc.styles['Heading 3'])
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    set_spacing_lines(p, 30, 10)
    set_run_font(p.add_run(text), ea=HEI, size=Pt(12), bold=True)
    return p


def para(text, indent=True, center=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(text), ea=SONG, size=Pt(12))
    return p


def blank(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        set_run_font(p.add_run(''), size=Pt(12))


def page_break():
    doc.add_page_break()


# ============================================================
# 封面
# ============================================================
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.5)
set_run_font(p.add_run('学    号：' + '_' * 24 + '　　年    级：2024级'), ea=SONG, size=Pt(14))

blank(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run('高级系统编程技术课程报告'), ea=HEI, size=Pt(26), bold=False)

blank(3)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
set_run_font(p.add_run('基于Rust的可持久化网络键值存储系统\n的设计与实现'), ea=HEI, size=Pt(22), bold=False)

blank(4)

cover_items = [
    ('专    业', '计算机科学与技术'),
    ('姓    名', '_' * 24),
    ('指导教师', '_' * 24),
    ('评 阅 人', '_' * 24),
]
for label, value in cover_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    set_run_font(p.add_run(label + '　　'), ea=SONG, size=Pt(16))
    set_run_font(p.add_run(value), ea=SONG, size=Pt(16))

blank(3)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run('2026年9月'), ea=SONG, size=Pt(15))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run('中国   南京'), ea=SONG, size=Pt(15))

page_break()

# ============================================================
# 目录
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run('目    录'), ea=HEI, size=Pt(18))
set_spacing_lines(p, 80, 50)

p = doc.add_paragraph()
run = p.add_run()
set_run_font(run, size=Pt(12))
f1, f2, f3, f4 = (OxmlElement('w:fldChar'), OxmlElement('w:instrText'),
                  OxmlElement('w:fldChar'), OxmlElement('w:fldChar'))
f1.set(qn('w:fldCharType'), 'begin')
f2.set(qn('xml:space'), 'preserve')
f2.text = ' TOC \\o "1-3" \\h \\z \\u '
f3.set(qn('w:fldCharType'), 'separate')
f4.set(qn('w:fldCharType'), 'end')
run._element.append(f1)
run._element.append(f2)
run._element.append(f3)
tip = p.add_run('（目录为自动生成域：请在 Word 中全选后按 F9，或右键选择“更新域”以生成目录）')
set_run_font(tip, size=Pt(12))
run._element.append(f4)

page_break()

# ============================================================
# 正文内容
# ============================================================
CONTENT = [
    ('c', '第1章 课程设计任务'),
    ('h2', '1.1 课程设计目标'),
    ('p', '本课程设计的目标是使用 Rust 实现一个支持数据持久化和多客户端访问的网络键值存储系统。系统由服务器端和客户端两部分组成：用户通过客户端可以完成数据写入、查询、覆盖、删除、列出键以及查看服务器状态等操作；服务器负责统一维护共享数据、处理多个客户端请求，并保证服务器重启以后，已经确认写入的数据仍然能够恢复。'),
    ('p', '系统的基础数据链路为：客户端命令输入 → 请求协议解析 → TCP 网络传输 → 服务器请求分发 → 内存 KV 存储 → WAL 持久化。'),
    ('p', '基础版本支持 set、get、delete、keys、status、ping、quit 七类命令，同时对非法命令、非法 JSON、非法 UTF-8、Key 不存在、请求参数错误、超长帧等常见异常情况进行了处理。'),
    ('p', '需要说明的是，本课程设计并不打算实现一个完整的工业级数据库。本项目主要研究的内容包括：Rust 基础机制、TCP 网络通信、应用层协议设计、多客户端并发访问、共享状态同步、WAL 持久化、崩溃恢复以及性能测试与实验分析。因此，分布式一致性、集群复制、事务系统、权限认证、跨节点容灾等内容不属于本次课程设计的实现范围。'),
    ('h2', '1.2 本人承担内容'),
    ('p', '在本项目中，我主要承担 Leader 的角色。相比单纯负责某一个模块，我的工作更多集中在总体方向、模块接口、系统集成、测试验收和最终演示设计上。主要工作包括以下几个方面：'),
    ('p', '（1）借助 AI 编程工具完成第一版能够运行的 MVP，并以实际运行起来的系统为对象，反向学习 Rust、TCP、并发和持久化相关知识；'),
    ('p', '（2）在第一版基础上重新梳理总体架构，推动客户端、协议、网络、存储和持久化模块职责分离；'),
    ('p', '（3）组织组内分工，提前约定 Request / Response、Store API、错误码以及模块调用方向等数据契约，减少网络层和数据层之间的直接耦合；'),
    ('p', '（4）针对基础实现进一步提出网络异步化、锁策略、并发模型和持久化性能等研究方向；'),
    ('p', '（5）负责最终答辩演示方案设计，将原本比较零散的系统功能收敛为 CRUD、Concurrency、Crash Recovery 和 Performance Lab 四个核心演示场景；'),
    ('p', '（6）设计性能实验室中的控制变量实验方法、前端原型以及前后端实验接口，并负责最终联调与整体演示流程。'),
    ('p', '项目整体并不是按照“先把所有理论学完，再开始写程序”的方式完成的，而是采用“先做出 MVP → 观察真实系统 → 针对实际问题补知识 → 重新分析设计 → 继续优化和验证”的迭代方式推进。'),

    ('c', '第2章 系统方案设计及创新性设计'),
    ('h2', '2.1 系统总体架构'),
    ('p', '系统按照职责划分为五个主要部分，自上而下依次为：命令行客户端 → 协议层 → 网络服务层 → 存储层 → 持久化层。各模块主要职责如下：'),
    ('p', '（1）客户端：负责用户输入、构造请求、发送请求并展示服务器返回结果。'),
    ('p', '（2）协议层：负责 Request / Response 数据结构、JSON 编解码、消息边界以及输入合法性检查。'),
    ('p', '（3）网络服务层：负责 TCP 监听、客户端连接管理、请求读取、请求分发以及响应发送。'),
    ('p', '（4）存储层：负责维护服务器当前的内存 KV 状态，并提供统一的 Store API。'),
    ('p', '（5）持久化层：负责将修改操作写入 WAL，并在服务器重新启动以后通过 WAL Replay 恢复数据。'),
    ('p', '这种分层并不是单纯为了让目录结构看起来更整齐，而是为了让不同部分能够相对独立地开发和测试。例如，Store 可以完全脱离 TCP 测试 CRUD；Protocol 可以单独验证 JSON 解析和分帧；Persistence 可以通过临时 WAL 文件测试崩溃恢复。这样一来，某一层发生修改时，也不需要把整个系统全部重写。'),
    ('h2', '2.2 网络协议设计'),
    ('p', '基础系统采用“TCP + UTF-8 + JSON Lines”作为网络协议。例如一个写入请求可以表示为 {"cmd":"set","key":"name","value":"Alice"}，每一条 Request 和 Response 都使用换行符“\\n”作为逻辑消息结束标志。'),
    ('p', '之所以需要自己定义消息边界，是因为 TCP 本身提供的只是可靠、有序的字节流，它并不知道什么叫“一条 JSON 请求”。客户端发送一次 send() 并不代表服务器就一定会通过一次 read() 得到完整请求：一个请求可能被拆成多次读取，多条请求也可能一起进入缓冲区。因此系统通过“缓冲读取 → 查找 LF → 提取完整一行 → JSON 解析”的流程解决 TCP 分段和粘包问题。'),
    ('p', '同时，系统对单条逻辑帧设置了 64 KiB 上限。这样即使客户端持续发送异常数据却一直不发送换行符，服务器也不会无限扩大缓冲区。'),
    ('h2', '2.3 内存存储设计'),
    ('p', '服务器的当前状态使用 BTreeMap<String, String> 维护。选择 BTreeMap 的一个直接原因是系统存在 keys 命令，它要求返回当前所有 Key 并按照字典序展示。如果使用 HashMap，还需要额外排序；而 BTreeMap 本身就按照 Key 有序组织，因此可以直接顺序遍历。对于当前课程设计的数据规模来说，BTreeMap 的性能完全能够满足要求，同时代码结构也比较直接。'),
    ('h2', '2.4 WAL 持久化设计'),
    ('p', '如果服务器只使用内存保存数据，那么服务器进程一旦退出，内存中的数据就会直接消失。因此系统引入 WAL（Write-Ahead Log）。形如 {"op":"set","key":"name","value":"Alice"}、{"op":"delete","key":"name"} 的记录会被顺序追加到 WAL 文件中。每次数据修改并不是重新把整个数据库写入磁盘，而是将本次操作追加到 WAL。'),
    ('p', '写操作按照“构造 WAL Record → append → flush → sync_data → 更新内存 → 返回成功”的顺序执行。这个顺序非常重要：如果先修改内存并向客户端返回成功、之后才写磁盘，那么服务器可能出现“客户端已经收到 OK → 服务器突然崩溃 → 磁盘还没有记录本次修改 → 服务器重启 → 刚才成功的数据消失”的情况。因此本系统要求：客户端收到成功响应之前，服务器必须已经拥有可以用于重启恢复的持久化依据。'),
    ('p', '服务器启动时，则执行“读取 WAL → 逐条 Replay → 重新执行 set / delete → 恢复 BTreeMap”的流程，从而重建服务器状态。'),
    ('h2', '2.5 并发安全设计'),
    ('p', '基础版本采用 Arc<Mutex<PersistentStore>> 管理共享 Store。其中 Arc 解决多个线程共同持有 Store 的问题；Mutex 负责保证同一时间只有一个线程进入需要互斥访问的临界区。'),
    ('p', '一个需要特别注意的地方是：Store 锁不能从接收到请求开始一直持有到响应发送结束。因此，网络等待、JSON 解析、请求构造、响应发送等操作都尽量不持有 Store 锁，只有真正需要访问共享数据时才进入临界区。对于写操作，WAL append、flush / sync_data 与内存更新需要处于统一的写入逻辑中，这样可以避免两个客户端同时修改数据时，出现 WAL 顺序和内存提交顺序不一致的问题。'),
    ('h2', '2.6 错误处理设计'),
    ('p', '系统没有采用“出错以后直接让 Server 崩溃”的方式处理客户端输入，而是尽量将可预期错误转换成结构化响应。错误大致可以分成四类：'),
    ('p', '（1）协议错误：如非法 JSON、非法 UTF-8、帧过长。这类问题发生在请求还没有真正进入业务逻辑以前。'),
    ('p', '（2）命令错误：如未知命令、参数缺失、参数格式错误。'),
    ('p', '（3）业务错误：如查询不存在的 Key、删除不存在的 Key。'),
    ('p', '（4）系统和持久化错误：如 WAL 打开失败、WAL 写入失败、磁盘同步失败、WAL Replay 异常。'),
    ('p', '其中客户端输入造成的错误应尽量转换成明确的 Response，而不是影响整个服务器；对于真正的 IO 和持久化错误，则通过 Rust 的错误返回机制向上层传播，再由服务器统一决定如何处理。这种设计的目标是：一个客户端发错请求，不应该把其他客户端正在使用的服务器一起带崩。'),
    ('h2', '2.7 可扩展性与可维护性设计'),
    ('p', '项目后期比较重视的一点，是减少不同模块之间互相知道太多内部实现。团队提前约定了 Request / Response、Store API、错误码、协议格式以及模块调用方向。例如网络层只需要知道 Store 能执行什么操作，而不应该直接知道 Store 内部到底使用 Mutex 还是 RwLock、WAL 文件内部怎么组织。这样以后即使将 Mutex 替换成 RwLock，或者将同步网络模型改成 Tokio，协议层和客户端也不需要大范围跟着修改。这也是后续能够进行多种性能方案对比的基础。'),
    ('h2', '2.8 并发模型与锁策略实验化'),
    ('p', '项目最初采用的是比较容易保证正确性的方案：同步的 Thread-per-connection 模型加 Mutex。后续没有直接把原实现删掉，而是保留不同实现，让它们成为性能实验变量。目前主要比较的并发模型是 Sync（Thread-per-connection）与 Async（Tokio），锁策略是 Mutex 与 RwLock。'),
    ('p', '这里一个比较重要的设计思想是：Async 不代表一定比 Sync 快，RwLock 也不代表一定比 Mutex 快，这些只能算理论上的性能假设。最终是否更快，还需要结合客户端数量、读写比例、锁竞争、网络等待、持久化开销和数据规模实际测试。因此项目没有把 Async 和 RwLock 简单写成“性能优化完成”，而是把它们变成了可以被实验验证的研究变量。'),
    ('h2', '2.9 Performance Lab 性能实验室'),
    ('p', '原来的 Performance 页面后来被重新设计成“性能实验室”。它不再只是生成几组简单性能数字，而是按照一个完整实验流程执行：固定实验条件 → 选择研究变量 → 运行对照组 → 恢复一致环境 → 运行实验组 → 收集数据 → 可视化比较 → 根据结果得出结论。'),
    ('p', '目前主要研究三个变量：一是并发模型（Sync / Async）；二是锁策略（Mutex / RwLock）；三是工作负载，其中 Read Heavy 为 90% 读 10% 写，Mixed 为 50% 读 50% 写，Write Heavy 为 10% 读 90% 写。客户端规模按照 1、10、50、100 Clients 逐级运行。每组实验统计 Throughput、P50、P95、P99、Success 和 Failed。其中 P99 主要用于观察高并发环境中的尾延迟：平均延迟有时候看起来并不高，但如果少数请求耗时特别长，用户实际体验仍然可能很差，因此 P95 / P99 是本次实验比较关注的数据。'),
    ('h2', '2.10 控制变量实验设计'),
    ('p', '为了让性能数据真正具有可解释性，系统会明确展示实验条件，包括 Dataset Size、Value Size、Requests / Scale、Persistence Strategy、Protocol 和 Network Environment。一组典型实验条件为：数据集 10,000 Keys，Value 大小 128 B，每档 10,000 个请求，持久化策略为 WAL + sync_data，协议为 JSON Lines，网络环境为 Localhost。'),
    ('p', '其中很重要的一点是：比较 Sync / Async 或 Mutex / RwLock 时，不能偷偷改变 WAL 和 sync_data 等持久化语义。例如如果 Sync 版本每次写入都执行 sync_data，而 Async 版本完全不做磁盘同步，那么即使后者快很多，也不能证明 Async 更快，因为两边实际上已经不是同一个实验条件。因此最终形成三类标准实验：Experiment A（Sync vs Async）、Experiment B（Mutex vs RwLock）、Experiment C（Read Heavy vs Mixed vs Write Heavy）。'),
    ('h2', '2.11 Crash Recovery 崩溃恢复实验'),
    ('p', '崩溃恢复演示也不再只是简单观察“服务器重启以后数据好像还在”，而是设计成一个自动一致性验证实验。系统明确区分三种状态：Memory Store、WAL 和 Frontend Verification Snapshot。'),
    ('p', '演示流程为：写入种子数据（Seed Data）→ 记录 Before 状态 → 强制终止 Server → Memory Store 消失而 WAL 文件仍然存在 → 重新启动 Server → WAL Replay → 重新构造 Memory Store → Before / After 自动比较。'),
    ('p', '这里需要特别说明：前端保存的 Before Snapshot 只负责验证恢复前后状态是否一致，它并不是服务器恢复数据的来源。真正的数据恢复链路始终是 WAL → Replay → BTreeMap。最终页面通过 Before Keys、Recovered Keys、Lost Keys、Before Hash、After Hash、WAL Replay Count、Recovery Time 等指标，生成 CONSISTENCY PASS 或 CONSISTENCY FAIL 的判定。这样原本依赖人工观察的恢复过程，就变成了一个可以自动判断结果的数据一致性实验。'),

    ('c', '第3章 方案论证及可行性分析'),
    ('h2', '3.1 网络协议方案可行性'),
    ('p', '本系统没有采用复杂的二进制协议，而是选择“TCP + JSON Lines”。主要原因是课程设计首先需要保证能够实现、容易调试、容易验证、能够解释。TCP 可以提供可靠、有序的数据传输，非常适合客户端和 KV Server 之间的请求响应模型。JSON 的可读性也比较强，例如直接看到 {"cmd":"get","key":"name"} 就能够判断客户端正在执行什么操作，这对于课程设计阶段的调试非常方便。JSON Lines 又解决了 TCP 本身没有消息边界的问题。虽然二进制协议在空间和性能上可能更加优秀，但对于本项目的数据规模和研究目标来说，JSON Lines 的开发成本更低，也更容易进行错误测试。因此该协议方案具有较好的实现可行性。'),
    ('h2', '3.2 存储方案可行性'),
    ('p', '系统内存层采用 BTreeMap。对于当前课程设计的数据规模，CRUD 的复杂度和内存消耗完全能够满足需求。同时 BTreeMap 天然维护 Key 顺序，可以简化 keys 命令的实现。本项目的重点并不是自行实现 B+ 树、LSM Tree 或数据库索引系统，因此直接使用 Rust 标准数据结构，可以把精力放到网络、并发和持久化问题上。这是一个在项目复杂度和功能目标之间比较合适的取舍。'),
    ('h2', '3.3 并发方案可行性'),
    ('p', '基础版本采用 Thread-per-connection 加 Arc<Mutex<PersistentStore>> 的结构。这种结构并不是理论上最高性能的网络服务器方案，但优势是模型简单、容易理解、容易验证、共享状态规则清晰，对于课程设计级别的连接数量已经能够正常工作。后续再加入 Async / Tokio 和 RwLock 进行性能研究。这样可以保证项目首先拥有一个正确并且稳定的基线版本，再讨论性能，这比一开始直接引入复杂异步结构更容易定位问题。'),
    ('h2', '3.4 数据持久化方案可行性'),
    ('p', '如果每执行一次 set 或 delete 就重新将整个数据库写入磁盘，那么随着数据量增大，写放大会越来越明显。WAL 则只需要追加本次修改（append record），实现相对简单，同时也能满足崩溃恢复需求。通过“append → flush → sync_data → Memory Update → ACK”的顺序，保证客户端收到成功响应之前已经存在恢复依据。重启以后通过“WAL Replay → 恢复 Memory Store”也能够比较直接地验证。因此 WAL 很适合本课程设计的数据持久化需求。'),
    ('h2', '3.5 性能优化方案可行性'),
    ('p', '理论上，Async 可以减少大量连接等待时每连接一个 OS Thread 带来的资源和调度成本。但是，网络并发优化并不等于存储吞吐无限增长，因为真正的写请求最终仍然需要经过共享 Store、WAL、flush 和 sync_data。如果瓶颈出现在磁盘持久化或者写锁竞争，那么即使网络层改成 Async，性能也不会无限提高。同样，RwLock 允许多个 Reader 同时访问，在 Read Heavy 场景下理论上可能比 Mutex 更有优势；但在 Write Heavy 环境中，大量请求仍然需要独占写锁，所以 RwLock 是否更快需要由实际读写比例决定。因此本项目采用控制变量实验，而不是直接根据理论给出性能结论。'),
    ('h2', '3.6 测试方案可行性'),
    ('p', '基础系统已经建立自动化测试，覆盖 CRUD、协议解析、TCP 分段、TCP 粘包、非法 JSON、非法 UTF-8、超长帧、多客户端访问、WAL 重启恢复、损坏 WAL 等场景。整个测试体系可以分成三层理解：一是功能正确性，通过自动化测试验证系统能不能正确完成操作；二是故障正确性，通过 Crash Recovery 验证进程异常退出以后数据还能不能恢复；三是性能行为，通过 Performance Lab 验证不同实现方式到底在什么条件下更快。三类测试分别关注 Correctness、Recovery 和 Performance，因此测试方案能够比较完整地覆盖系统当前的主要设计目标。'),

    ('c', '第4章 设计过程及设计成果'),
    ('h2', '4.1 团队分工'),
    ('p', '本项目采用按模块职责进行分工，而不是简单按照代码文件数量平均分配。团队分工如表4.1所示。'),
    ('table4.1', None),
    ('p', '团队协作时首先确定公共数据契约，例如 Request / Response、Store API、错误码、协议格式和模块依赖方向，然后不同成员围绕接口分别开发。这样可以减少最后集成阶段出现“每个人自己的模块都能跑，但是拼起来接口完全对不上”的情况。'),
    ('h2', '4.2 项目进度安排'),
    ('p', '整个项目大致分为六个阶段：'),
    ('p', '（1）第一阶段：需求分析与 MVP。分析课程要求、确定系统范围、完成最小可运行版本，打通 Client → Server → CRUD → WAL 的完整链路。'),
    ('p', '（2）第二阶段：架构重新梳理。学习已有 MVP、重新划分模块，确定 Request / Response 与 Store API，制定团队接口。'),
    ('p', '（3）第三阶段：分模块开发与第一次集成。完成协议模块、网络模块、Store、WAL、客户端的开发与模块测试，并进行第一次系统集成。'),
    ('p', '（4）第四阶段：异常和边界测试。重点处理粘包、拆包、非法输入、多客户端、WAL 恢复与损坏 WAL 等场景。'),
    ('p', '（5）第五阶段：性能与崩溃恢复实验。完成 Sync / Async、Mutex / RwLock 对比、Benchmark Harness、Performance Lab 以及 Crash Recovery 自动验证。'),
    ('p', '（6）第六阶段：最终联调与答辩。完成前后端联调、测试回归、准备演示数据并组织答辩流程。'),
    ('h2', '4.3 程序设计与开发流程'),
    ('p', '整个项目的实际演进过程可以概括为：分析课程要求 → 完成第一版 MVP → 结合 MVP 学习 Rust 与系统知识 → 重新梳理模块边界 → 制定 Request / Response 与 Store API → 团队分模块实现与测试 → 第一次系统集成 → 异常与边界测试 → 基础系统验收 → 分析性能瓶颈 → 提出并发模型与锁策略研究方向 → 实现多个可比较方案 → 设计控制变量实验 → 开发 Performance Lab → 加入 Crash Recovery 自动一致性验证 → 前后端联调 → 形成最终答辩流程。'),
    ('p', '整个过程比较明显地经历了三个阶段：先做出来，理解为什么这样做，进一步验证怎样做更好。'),
    ('h2', '4.4 调试过程'),
    ('p', '课程设计过程中，一些问题也是在系统真正运行以后才暴露出来的。'),
    ('h3', '（1）一次 TCP Read 不等于一次 Request'),
    ('p', '早期如果直接认为一次 read 就等于一条 JSON，系统在简单测试中可能没有问题，但实际 TCP 可能出现一个请求被拆开，或者多个请求一次收到的情况。因此后来通过 Buffer 加 LF 帧边界重新实现消息分帧，并加入对应的拆包和粘包测试。'),
    ('h3', '（2）多个客户端共享同一个 Store'),
    ('p', '当多个客户端同时访问服务器时，Store 不再属于单一线程，于是需要解决谁拥有 Store、多个线程如何访问、写入时如何避免状态冲突等问题。最终通过 Arc 加 Mutex 管理共享状态，同时进一步限制锁的持有范围，避免网络等待期间长期占用 Store。'),
    ('h3', '（3）内存写入成功不代表数据真正安全'),
    ('p', '最开始只考虑 set 成功以后 BTreeMap 已经变化，后来加入崩溃测试以后才真正意识到：内存里的成功和客户端能够相信的“持久化成功”不是一回事。因此重新确认 WAL → flush → sync_data → Memory → ACK 的顺序，并通过重启恢复进行验证。'),
    ('h3', '（4）理论上的性能优化无法直接下结论'),
    ('p', '在引入 Async 和 RwLock 以后，最开始很容易直接认为新方案一定更好。但实际分析发现，性能可能同时受到网络、线程调度、锁竞争、WAL、sync_data 和读写比例的影响。因此后来没有单独比较某一次运行结果，而是建立 Performance Lab，固定实验变量后再进行 A/B 对照。'),
    ('h2', '4.5 最终设计成果'),
    ('p', '目前项目已经形成以下主要成果：Rust KV Server、命令行 Client、JSON Lines 网络协议、BTreeMap 内存 Store、WAL 持久化机制、多客户端访问、共享状态同步、自动化测试体系、Crash Recovery 实验、Performance Lab 以及 Web 演示界面。基础版本已经完成 41 个自动化测试。'),
    ('p', '最终演示内容被收敛为 CRUD → Concurrency → Crash Recovery → Performance Lab 四个部分。这个顺序实际上也对应系统分析的四个层次：功能能不能工作，多个客户端一起工作是否正确，服务器崩溃以后是否仍然正确，系统在不同条件下性能表现如何。'),
    ('h2', '4.6 最终演示效果'),
    ('p', '（1）CRUD：通过 Web 页面或客户端直接执行 SET、GET、UPDATE、DELETE、KEYS，展示服务器基础功能。'),
    ('p', '（2）Concurrency：启动多个并发客户端请求，验证服务器在共享 Store 环境下仍然能够正确处理数据。'),
    ('p', '（3）Crash Recovery：执行“写入数据 → 记录 Before → Kill Server → Restart → WAL Replay → Compare”的流程，页面最终输出 CONSISTENCY PASS / FAIL，用于证明恢复前后的数据一致性。'),
    ('p', '（4）Performance Lab：分别对 Sync / Async、Mutex / RwLock、不同 Workload 和不同 Client Scale 进行实验，最终展示 Throughput、P50、P95、P99、Success Rate 等数据。'),

    ('c', '第5章 工程管理与质量控制'),
    ('h2', '5.1 任务分解'),
    ('p', '在多人项目中，我逐渐认识到：真正重要的并不是把代码文件平均分给每个人，而是提前定义模块责任和数据接口。因此开发前优先确定 Request / Response、Store API、错误码、协议格式和模块调用方向，再根据这些接口拆分任务。这样协议模块、网络模块和存储模块可以相对独立开发。'),
    ('h2', '5.2 测试覆盖'),
    ('p', '项目基础版本目前拥有 41 个自动化测试，测试内容包括基础 CRUD、协议解析、TCP 分段、TCP 粘包、非法 JSON、非法 UTF-8、Key 不存在、帧大小限制、多客户端访问、WAL Replay、服务器重启恢复以及损坏 WAL。对于后续的 Performance Lab，则不再只测试程序有没有报错，而是增加实验条件、实验结果、公平性和结果可重复性方面的检查。'),
    ('h2', '5.3 代码规范'),
    ('p', '项目在开发过程中主要通过统一模块职责、统一 Request / Response、统一 Store API、统一错误返回方式、避免跨层直接访问内部状态、尽量减少无关重构等方式控制代码质量。Rust 代码遵循统一的命名和格式习惯，公共接口修改时需要同步相关模块，而不是让各成员分别维护互不兼容的版本；对于可能影响多个模块的数据结构修改，需要先明确接口，再进行实现。'),
    ('h2', '5.4 集成过程'),
    ('p', '项目第一次集成之后，一个比较明显的问题是：一个看似很小的功能修改，有时会穿过很多模块。例如增加 Lock = RwLock 这个实验变量，实际上可能影响前端实验配置、实验 API、网络运行模式、服务端配置、Benchmark Harness、Result 数据结构以及前端图表。个人开发时可以一次性全部修改，但多人协作时，如果没有统一接口，就很容易出现 A 等 B、B 等 C、C 使用旧字段、前端又已经改成新字段的情况。'),
    ('p', '因此后期集成逐渐采用“先确定数据契约 → 各模块实现 → 模块测试 → 集成 → 系统回归”的方式进行。'),
    ('h2', '5.5 进度与范围控制'),
    ('p', '项目进行过程中会不断产生新的想法。如果每出现一个 Idea 就马上加入当前版本，很容易导致项目不断扩大，最后核心功能反而没有收尾。因此后期开始将需求区分为当前版本必须完成、下一阶段完成、实验性功能和暂不实施四类。例如本轮 Performance Lab 明确决定不继续扩展无关 Dashboard、不加入 Pub/Sub，而是集中完成 CRUD、Concurrency、Crash Recovery 和 Performance Lab 四个答辩核心场景。这种范围控制让后期开发目标更加稳定。'),
    ('h2', '5.6 性能实验质量控制'),
    ('p', '性能实验还有一项不同于普通功能测试的要求：必须保证实验公平。例如比较两个方案时，只允许改变当前正在研究的变量，其他条件如 Dataset、Value Size、Request Count、Persistence、Protocol 和 Network 都尽量保持一致；实验 A/B 之间也需要尽可能恢复相同的数据初始状态。否则最后观察到的性能差异可能根本不是 Mutex 与 RwLock 造成的，而只是两次测试环境不同。因此，控制变量本身也是 Performance Lab 质量控制的一部分。'),

    ('c', '第6章 课程设计收获及体会'),
    ('h2', '6.1 Rust 学习收获'),
    ('p', '本次课程设计对我的一个明显帮助，是把之前比较抽象的 Rust 概念放到了实际系统中。项目开始前，我对所有权、借用、Arc、Mutex 等内容基本没有实际使用经验。如果只看语法，Arc<Mutex<T>> 很容易只是一个需要记忆的写法；但放到这个项目里以后，它对应的是一个很具体的问题：服务器有多个客户端线程，大家都需要访问同一个 Store，单一所有权已经不够，Arc 负责共享所有权，Mutex 负责控制并发访问。这样以后再看到这些概念，就不再只是 Rust 的语法规则，而是能够对应到真实系统里的共享状态问题。'),
    ('h2', '6.2 系统设计方面的收获'),
    ('p', '这次项目让我比较明显地认识到：程序能够工作只是第一层，发生故障以后仍然正确是第二层，知道系统为什么快或慢又是另外一层。最开始看到 WAL 时，我只是把它理解为“把数据保存到磁盘”，后来才逐渐理解它还涉及写入顺序、flush、sync_data、ACK 时机、崩溃恢复以及内存和磁盘一致性。TCP 也是类似，开始时很容易理解成“Client 把 JSON 发给 Server”，进一步学习以后才明白，TCP 只负责传输字节，并不知道 JSON 从哪里开始、到哪里结束，因此应用层还需要自己定义消息边界。这些问题只有放到一个完整系统里以后，才能真正连接起来理解。'),
    ('h2', '6.3 性能实验与问题解决方面的收获'),
    ('p', 'Performance Lab 是这次项目后期对我影响比较大的部分。最开始很容易有这样的直觉：“Async 比 Sync 高级，所以 Async 应该更快；RwLock 可以有多个 Reader，所以 RwLock 应该更快。”但实际系统并没有这么简单，性能同时受到并发量、工作负载、锁竞争、WAL、磁盘同步、网络等待和数据规模等条件的影响。'),
    ('p', '因此后来整个思路逐渐从“我觉得这个方案应该更快”，变成：提出一个假设 → 确定研究变量 → 固定其他条件 → 运行 A/B 实验 → 统计 Throughput 与 P50 / P95 / P99 → 观察数据 → 再得出结论。我觉得这比最后某个方案到底快了多少更加重要，因为它让我开始意识到：代码能够运行或者理论上合理，都不等于结论已经被证明。'),
    ('h2', '6.4 团队协作与工程管理方面的收获'),
    ('p', '在团队协作方面，我也逐渐认识到：Leader 比较困难的地方，不是把第一轮任务分下去，而是项目进行过程中不断会产生新的需求和接口变化。例如只是增加一个新的实验字段，就可能同时影响前端、API、Server、Benchmark、Result 和图表。如果每个成员都按照自己的理解修改，很容易在最后集成时产生大量额外沟通成本。'),
    ('p', '因此相比项目开始时，我现在更认可“先确定版本目标和接口，再进行开发”。项目进行过程中也应该尽量减少没有边界的新需求；如果确实需要修改公共接口，则应该先和相关成员同步，而不是一个模块先改完以后，再让其他人被动适配。'),
    ('h2', '6.5 总结'),
    ('p', '整个课程设计的学习路径最终可以概括为：先让程序运行 → 理解系统为什么这样运行 → 重新划分模块责任 → 保证功能正确 → 验证故障情况下是否正确 → 提出性能假设 → 设计实验 → 收集数据 → 根据证据判断优化是否有效。'),
    ('p', '本次课程设计最终完成的不只是一个 Rust 网络 KV 系统。对我来说，更重要的变化是从最开始关注“功能有没有实现”，逐渐转向思考系统为什么这样设计、如果崩溃会发生什么、多人同时访问是否安全、性能瓶颈在哪里、怎样设计实验证明自己的判断。相比单纯增加了多少功能，我认为这种从功能实现逐渐转向系统分析、工程验证和实验思维的过程，是本次课程设计最有价值的收获。'),

    ('c', '参考文献'),
    ('ref', '[1] Klabnik S, Nichols C. The Rust Programming Language[M]. 2nd ed. San Francisco: No Starch Press, 2023.'),
    ('ref', '[2] Stevens W R, Fenner B, Rudoff A M. UNIX Network Programming, Volume 1: The Sockets Networking API[M]. 3rd ed. Boston: Addison-Wesley, 2004.'),
    ('ref', '[3] Bray T. The JavaScript Object Notation (JSON) Data Interchange Format: RFC 8259[S]. IETF, 2017.'),
    ('ref', '[4] Tokio Project. Tokio: An asynchronous runtime for the Rust language[EB/OL]. https://tokio.rs, 2025.'),
    ('ref', '[5] The Rust Documentation. std::collections::BTreeMap[EB/OL]. https://doc.rust-lang.org/std/collections/struct.BTreeMap.html, 2025.'),
]

TEAM_TABLE = [
    ('成员', '主要工作'),
    ('许赵泓 / Leader', '总体架构、接口协调、任务拆分、系统集成、Performance Lab、测试验收、答辩设计'),
    ('张良俊', '网络层：TCP Server、连接处理、多客户端、线程、共享状态、异步改造'),
    ('姚必顺', '数据层：BTreeMap、CRUD、WAL、恢复、异常数据、高并发相关测试'),
]

for kind, text in CONTENT:
    if kind == 'c':
        chapter(text)
    elif kind == 'h2':
        h2(text)
    elif kind == 'h3':
        h3(text)
    elif kind == 'p':
        para(text)
    elif kind == 'ref':
        para(text)
    elif kind == 'table4.1':
        # 表题
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.line_spacing = 1.5
        set_run_font(cap.add_run('表4.1 项目团队分工'), ea=HEI, size=Pt(10.5), bold=True)
        table = doc.add_table(rows=len(TEAM_TABLE), cols=2)
        table.style = doc.styles['Table Grid']
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (a, b) in enumerate(TEAM_TABLE):
            for j, val in enumerate((a, b)):
                cell = table.cell(i, j)
                cp = cell.paragraphs[0]
                cp.paragraph_format.line_spacing = 1.5
                if i == 0:
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run_font(cp.add_run(val), ea=SONG, size=Pt(10.5), bold=(i == 0))
        doc.add_paragraph()

# 参考文献标题补丁：把最后一个 'c' 的样式改为左对齐黑体小二（参考文献一级标题）
for p in doc.paragraphs:
    if p.text == '参考文献':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        break

OUT = r'f:\college\junior\这里的水很深\Rust课设\高级系统编程技术课程报告-基于Rust的可持久化网络键值存储系统.docx'
doc.save(OUT)
print('saved:', OUT)
